#!/usr/bin/env python3
"""Build the public DOCX specification from the canonical Markdown source."""

from __future__ import annotations

import argparse
import hashlib
import re
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "rls_asic_spec.md"
OUTPUT = ROOT / "docs" / "rls_asic_spec.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "526175"
TABLE_FILL = "E8EEF5"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, size: float | None = None, bold: bool | None = None,
             color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_style(style, size: float, color: str, before: float, after: float,
                    line_spacing: float, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def table_widths(columns: int) -> list[int]:
    patterns = {
        3: [2700, 1850, 4810],
        4: [2580, 1080, 1320, 4380],
    }
    return patterns.get(columns, [CONTENT_DXA // columns] * (columns - 1) + [CONTENT_DXA - (CONTENT_DXA // columns) * (columns - 1)])


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = table_widths(len(table.columns))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            row_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_pr.append(repeat)
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), TABLE_FILL)
                cell._tc.get_or_add_tcPr().append(shade)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_font(run, 9.0, bold=(row_index == 0), color="172033")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, display, end))
    set_font(run, 9, color=MUTED)


def parse_inline(paragraph, value: str, size: float | None = None, color: str | None = None) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", value)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=size or 10.5, color=color, bold=False)
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, color=color)


def build(markdown: str, digest: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(doc.styles["Normal"], 11, "172033", 0, 6, 1.25)
    configure_style(doc.styles["Heading 1"], 16, BLUE, 18, 10, 1.0, True)
    configure_style(doc.styles["Heading 2"], 13, BLUE, 14, 7, 1.0, True)
    configure_style(doc.styles["Heading 3"], 12, DARK_BLUE, 10, 5, 1.0, True)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    set_font(header_p.add_run("RLS ASIC/FPGA Portable RTL Specification"), 9, color=MUTED, bold=True)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer_p.add_run("Public evidence reference | Page "), 9, color=MUTED)
    add_field(footer_p, "PAGE")

    lines = markdown.splitlines()
    title_text = lines[0].removeprefix("# ")
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(5)
    parse_inline(title, title_text, 23, "172033")
    title.runs[0].bold = True
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    parse_inline(subtitle, "12-tap complex RLS self-interference cancellation Soft IP", 12, MUTED)
    fingerprint = doc.add_paragraph()
    fingerprint.paragraph_format.space_after = Pt(14)
    parse_inline(fingerprint, f"Markdown SHA-256: {digest}", 8.5, MUTED)

    index = 1
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("**"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            parse_inline(paragraph, line[:-1] if line.endswith("\\") else line, 9.5, MUTED)
            index += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            index += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            if len(rows) > 1 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
                rows.pop(1)
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for row_index, values in enumerate(rows):
                for col_index, value in enumerate(values):
                    cell = table.cell(row_index, col_index)
                    cell.text = ""
                    parse_inline(cell.paragraphs[0], value)
            format_table(table)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(1)
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            parse_inline(paragraph, line[2:])
            index += 1
            continue
        paragraph = doc.add_paragraph()
        parse_inline(paragraph, line.replace("  ", " "))
        index += 1

    doc.core_properties.title = title_text
    doc.core_properties.subject = f"Markdown-SHA256:{digest}"
    doc.core_properties.keywords = f"Markdown-SHA256:{digest}"
    doc.core_properties.creator = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    return doc


def scrub_ooxml(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as source, tempfile.NamedTemporaryFile(
        delete=False, suffix=".docx", dir=path.parent
    ) as temp:
        temp_path = Path(temp.name)
        with zipfile.ZipFile(temp, "w", compression=zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                if info.filename == "docProps/custom.xml":
                    continue
                data = source.read(info.filename)
                if info.filename.endswith(".xml"):
                    try:
                        root = ET.fromstring(data)
                        for element in list(root.iter()):
                            if element.tag.rsplit("}", 1)[-1] in {"creator", "lastModifiedBy"}:
                                element.text = None
                            for attribute in list(element.attrib):
                                if attribute.rsplit("}", 1)[-1].startswith("rsid"):
                                    del element.attrib[attribute]
                            for child in list(element):
                                if child.tag.rsplit("}", 1)[-1].startswith("rsid"):
                                    element.remove(child)
                        data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                    except ET.ParseError:
                        pass
                target.writestr(info, data)
    temp_path.replace(path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    markdown_bytes = args.source.read_bytes()
    digest = hashlib.sha256(markdown_bytes).hexdigest()
    document = build(markdown_bytes.decode("utf-8"), digest)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    scrub_ooxml(args.output)
    print(f"SPEC_DOCX_GENERATED markdown_sha256={digest}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
